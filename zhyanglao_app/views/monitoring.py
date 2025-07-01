from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from ..models import Elder, Device, HealthData, Alert
from django.utils import timezone
from datetime import timedelta
import random  # 临时用于模拟数据
import xlsxwriter # type: ignore
import io
from docx import Document # type: ignore
from docx.shared import Inches # type: ignore
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
import json

@login_required
def monitoring_dashboard(request):
    """实时监控仪表盘"""
    elders = Elder.objects.all()
    devices = Device.objects.all()
    alerts = Alert.objects.filter(
        timestamp__gte=timezone.now() - timedelta(days=1)
    ).order_by('-timestamp')[:10]
    
    context = {
        'elders': elders,
        'devices': devices,
        'alerts': alerts,
    }
    return render(request, 'monitoring/dashboard.html', context)

@login_required
def elder_monitoring(request, elder_id):
    """老人监控详情"""
    elder = get_object_or_404(Elder, id=elder_id)
    devices = Device.objects.filter(elder=elder)
    health_data = HealthData.objects.filter(
        elder=elder,
        timestamp__gte=timezone.now() - timedelta(days=1)
    ).order_by('-timestamp')
    
    context = {
        'elder': elder,
        'devices': devices,
        'health_data': health_data,
    }
    return render(request, 'monitoring/elder_detail.html', context)

@login_required
def get_real_time_data(request, elder_id):
    """获取实时监控数据"""
    elder = get_object_or_404(Elder, id=elder_id)
    latest_health = HealthData.objects.filter(elder=elder).order_by('-timestamp').first()
    devices = Device.objects.filter(elder=elder)
    
    device_status = [{
        'id': device.device_id,
        'name': device.name,
        'type': device.device_type,
        'status': device.status,
        'last_active': device.last_active.isoformat() if device.last_active else None
    } for device in devices]
    
    data = {
        'elder_id': elder.id,
        'name': f'{elder.name}',
        'devices': device_status,
        'health_data': {
            'timestamp': latest_health.timestamp.isoformat() if latest_health else None,
            'heart_rate': latest_health.heart_rate if latest_health else None,
            'blood_pressure_sys': latest_health.blood_pressure_sys if latest_health else None,
            'blood_pressure_dia': latest_health.blood_pressure_dia if latest_health else None,
            'temperature': str(latest_health.temperature) if latest_health and latest_health.temperature else None,
            'sleep_quality': latest_health.sleep_quality if latest_health else None,
            'activity_level': latest_health.activity_level if latest_health else None
        } if latest_health else None
    }
    return JsonResponse(data)

@login_required
def alert_list(request):
    """告警列表"""
    alerts = Alert.objects.all().order_by('-timestamp')
    return render(request, 'monitoring/alert_list.html', {'alerts': alerts})

@login_required
def get_elder_health_data(request, elder_id):
    """获取老人健康数据API - 支持不同时间周期"""
    elder = get_object_or_404(Elder, id=elder_id)
    period = request.GET.get('period', 'week')
    
    # 根据周期确定查询时间范围
    if period == 'day':
        start_date = timezone.now() - timedelta(days=1)
    elif period == 'week':
        start_date = timezone.now() - timedelta(days=7)
    elif period == 'month':
        start_date = timezone.now() - timedelta(days=30)
    else:
        start_date = timezone.now() - timedelta(days=7)
    
    # 查询健康数据
    health_records = HealthData.objects.filter(
        elder=elder,
        timestamp__gte=start_date
    ).order_by('timestamp')
    
    # 如果没有数据，生成模拟数据
    if not health_records.exists():
        return generate_mock_health_data(elder, period)
    
    # 提取数据
    heart_rate = []
    systolic = []
    diastolic = []
    blood_oxygen = []
    temperature = []
    dates = []
    
    for record in health_records:
        heart_rate.append(record.heart_rate)
        systolic.append(record.systolic)
        diastolic.append(record.diastolic)
        blood_oxygen.append(record.blood_oxygen)
        temperature.append(float(record.temperature) if record.temperature else None)
        dates.append(record.timestamp.strftime('%m-%d'))
    
    # 获取最新数据
    latest = health_records.last()
    latest_data = {
        'heart_rate': latest.heart_rate,
        'systolic': latest.systolic,
        'diastolic': latest.diastolic,
        'blood_oxygen': latest.blood_oxygen,
        'temperature': float(latest.temperature) if latest.temperature else None,
        'steps': latest.steps,
        'sleep_hours': latest.sleep_hours,
        'timestamp': latest.timestamp.strftime('%Y-%m-%d %H:%M:%S')
    }
    
    return JsonResponse({
        'status': 'success',
        'data': {
            'heart_rate': heart_rate,
            'systolic': systolic,
            'diastolic': diastolic,
            'blood_oxygen': blood_oxygen,
            'temperature': temperature,
            'dates': dates
        },
        'latest': latest_data
    })

@login_required
def get_elder_latest_health_data(request, elder_id):
    """获取老人最新健康数据API"""
    elder = get_object_or_404(Elder, id=elder_id)
    
    # 查询最新健康数据
    latest = HealthData.objects.filter(elder=elder).order_by('-timestamp').first()
    
    if not latest:
        # 如果没有数据，生成模拟数据
        return generate_mock_latest_health_data(elder)
    
    data = {
        'heart_rate': latest.heart_rate,
        'systolic': latest.systolic,
        'diastolic': latest.diastolic,
        'blood_oxygen': latest.blood_oxygen,
        'temperature': float(latest.temperature) if latest.temperature else None,
        'steps': latest.steps,
        'sleep_hours': latest.sleep_hours,
        'timestamp': latest.timestamp.strftime('%Y-%m-%d %H:%M:%S')
    }
    
    return JsonResponse({
        'status': 'success',
        'data': data
    })

def generate_mock_health_data(elder, period):
    """生成模拟健康数据 - 仅在没有实际数据时使用"""
    today = timezone.now()
    
    if period == 'day':
        days = 1
        interval = 2  # 小时
    elif period == 'week':
        days = 7
        interval = 1  # 天
    else:  # month
        days = 30
        interval = 2  # 天
    
    heart_rate = []
    systolic = []
    diastolic = []
    blood_oxygen = []
    temperature = []
    dates = []
    
    # 创建数据点
    for i in range(days // interval):
        date = today - timedelta(days=days-i*interval)
        dates.append(date.strftime('%m-%d'))
        
        # 生成合理范围内的随机数值
        heart_rate.append(random.randint(65, 85))
        systolic.append(random.randint(110, 130))
        diastolic.append(random.randint(70, 85))
        blood_oxygen.append(random.randint(95, 99))
        temperature.append(round(36 + random.random() * 1.2, 1))
    
    # 最新数据
    latest_data = {
        'heart_rate': heart_rate[-1],
        'systolic': systolic[-1],
        'diastolic': diastolic[-1],
        'blood_oxygen': blood_oxygen[-1],
        'temperature': temperature[-1],
        'steps': random.randint(3000, 8000),
        'sleep_hours': round(random.uniform(6, 8.5), 1),
        'timestamp': today.strftime('%Y-%m-%d %H:%M:%S')
    }
    
    return JsonResponse({
        'status': 'success',
        'data': {
            'heart_rate': heart_rate,
            'systolic': systolic,
            'diastolic': diastolic,
            'blood_oxygen': blood_oxygen,
            'temperature': temperature,
            'dates': dates
        },
        'latest': latest_data
    })

def generate_mock_latest_health_data(elder):
    """生成模拟的最新健康数据 - 仅在没有实际数据时使用"""
    data = {
        'heart_rate': random.randint(65, 85),
        'systolic': random.randint(110, 130),
        'diastolic': random.randint(70, 85),
        'blood_oxygen': random.randint(95, 99),
        'temperature': round(36 + random.random() * 1.2, 1),
        'steps': random.randint(3000, 8000),
        'sleep_hours': round(random.uniform(6, 8.5), 1),
        'timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    
    return JsonResponse({
        'status': 'success',
        'data': data
    })

@login_required
def export_health_data_excel(request, elder_id):
    """导出健康数据为Excel文件"""
    elder = get_object_or_404(Elder, id=elder_id)
    
    # 查询最近30天的健康数据
    start_date = timezone.now() - timedelta(days=30)
    health_records = HealthData.objects.filter(
        elder=elder,
        timestamp__gte=start_date
    ).order_by('timestamp')
    
    # 创建内存文件对象
    output = io.BytesIO()
    
    # 创建Excel工作簿和工作表
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet(f"{elder.name}的健康数据")
    
    # 添加标题
    bold = workbook.add_format({'bold': True})
    worksheet.write(0, 0, "日期", bold)
    worksheet.write(0, 1, "时间", bold)
    worksheet.write(0, 2, "心率(次/分)", bold)
    worksheet.write(0, 3, "收缩压(mmHg)", bold)
    worksheet.write(0, 4, "舒张压(mmHg)", bold)
    worksheet.write(0, 5, "血氧饱和度(%)", bold)
    worksheet.write(0, 6, "体温(°C)", bold)
    worksheet.write(0, 7, "步数", bold)
    worksheet.write(0, 8, "睡眠时长(小时)", bold)
    
    # 写入数据
    for i, record in enumerate(health_records):
        worksheet.write(i+1, 0, record.timestamp.strftime('%Y-%m-%d'))
        worksheet.write(i+1, 1, record.timestamp.strftime('%H:%M:%S'))
        worksheet.write(i+1, 2, record.heart_rate if record.heart_rate else "")
        worksheet.write(i+1, 3, record.systolic if record.systolic else "")
        worksheet.write(i+1, 4, record.diastolic if record.diastolic else "")
        worksheet.write(i+1, 5, record.blood_oxygen if record.blood_oxygen else "")
        worksheet.write(i+1, 6, float(record.temperature) if record.temperature else "")
        worksheet.write(i+1, 7, record.steps if record.steps else "")
        worksheet.write(i+1, 8, record.sleep_hours if record.sleep_hours else "")
    
    # 设置列宽
    worksheet.set_column(0, 0, 12)  # 日期列宽
    worksheet.set_column(1, 1, 10)  # 时间列宽
    worksheet.set_column(2, 8, 15)  # 其他列宽
    
    # 如果没有数据，添加示例数据（临时）
    if not health_records.exists():
        # 添加提示
        worksheet.merge_range(1, 0, 1, 8, "无健康数据记录，以下为示例数据", workbook.add_format({'italic': True}))
        
        # 添加7天示例数据
        for i in range(7):
            date = timezone.now() - timedelta(days=i)
            worksheet.write(i+2, 0, date.strftime('%Y-%m-%d'))
            worksheet.write(i+2, 1, "08:00:00")
            worksheet.write(i+2, 2, random.randint(65, 85))
            worksheet.write(i+2, 3, random.randint(110, 130))
            worksheet.write(i+2, 4, random.randint(70, 85))
            worksheet.write(i+2, 5, random.randint(95, 99))
            worksheet.write(i+2, 6, round(36 + random.random() * 1.2, 1))
            worksheet.write(i+2, 7, random.randint(3000, 8000))
            worksheet.write(i+2, 8, round(random.uniform(6, 8.5), 1))
    
    workbook.close()
    
    # 设置响应头
    output.seek(0)
    filename = f"{elder.name}健康数据_{timezone.now().strftime('%Y%m%d')}.xlsx"
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response

@login_required
def export_health_data_word(request, elder_id):
    """导出健康数据为Word文件"""
    elder = get_object_or_404(Elder, id=elder_id)
    
    # 查询最近30天的健康数据
    start_date = timezone.now() - timedelta(days=30)
    health_records = HealthData.objects.filter(
        elder=elder,
        timestamp__gte=start_date
    ).order_by('timestamp')
    
    # 创建Word文档
    document = Document()
    
    # 添加标题
    document.add_heading(f'{elder.name}的健康报告', 0)
    
    # 添加基本信息
    document.add_heading('基本信息', level=1)
    document.add_paragraph(f'姓名: {elder.name}')
    document.add_paragraph(f'年龄: {elder.age}岁')
    document.add_paragraph(f'性别: {"男" if elder.gender == "M" else "女" if elder.gender == "F" else "未知"}')
    document.add_paragraph(f'地址: {elder.address}')
    document.add_paragraph(f'联系电话: {elder.phone}')
    document.add_paragraph(f'紧急联系人: {elder.emergency_contact}')
    document.add_paragraph(f'紧急联系电话: {elder.emergency_phone}')
    document.add_paragraph(f'报告生成时间: {timezone.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # 添加健康数据
    document.add_heading('最近30天健康数据', level=1)
    
    # 如果没有数据
    if not health_records.exists():
        document.add_paragraph('暂无健康数据记录。', style='Intense Quote')
    else:
        # 添加数据表格
        table = document.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '日期'
        header_cells[1].text = '心率(次/分)'
        header_cells[2].text = '收缩压(mmHg)'
        header_cells[3].text = '舒张压(mmHg)'
        header_cells[4].text = '血氧饱和度(%)'
        header_cells[5].text = '体温(°C)'
        header_cells[6].text = '步数'
        header_cells[7].text = '睡眠时长(小时)'
        
        # 数据行
        for record in health_records:
            row_cells = table.add_row().cells
            row_cells[0].text = record.timestamp.strftime('%Y-%m-%d')
            row_cells[1].text = str(record.heart_rate) if record.heart_rate else "-"
            row_cells[2].text = str(record.systolic) if record.systolic else "-"
            row_cells[3].text = str(record.diastolic) if record.diastolic else "-"
            row_cells[4].text = str(record.blood_oxygen) if record.blood_oxygen else "-"
            row_cells[5].text = str(record.temperature) if record.temperature else "-"
            row_cells[6].text = str(record.steps) if record.steps else "-"
            row_cells[7].text = str(record.sleep_hours) if record.sleep_hours else "-"
    
    # 添加健康建议
    document.add_heading('健康建议', level=1)
    document.add_paragraph('根据您的健康数据，我们建议：')
    document.add_paragraph('1. 保持规律生活，早睡早起。', style='List Number')
    document.add_paragraph('2. 控制血压，减少高盐高脂食物摄入。', style='List Number')
    document.add_paragraph('3. 每天保持适量运动，建议步行30分钟以上。', style='List Number')
    document.add_paragraph('4. 定期测量血压、心率等指标，保持记录。', style='List Number')
    document.add_paragraph('5. 按时服药，不要随意停药或更改剂量。', style='List Number')
    
    # 保存到内存流
    output = io.BytesIO()
    document.save(output)
    
    # 设置响应头
    output.seek(0)
    filename = f"{elder.name}健康报告_{timezone.now().strftime('%Y%m%d')}.docx"
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    return response

def check_new_alerts(request):
    """API端点：检查自上次检查以来的新告警"""
    try:
        # 获取上次检查时间
        since_str = request.GET.get('since', '')
        if since_str:
            try:
                since = datetime.fromisoformat(since_str.replace('Z', '+00:00'))
            except ValueError:
                # 如果时间格式不正确，默认使用1分钟前
                since = timezone.now() - timedelta(minutes=1)
        else:
            # 默认检查1分钟内的告警
            since = timezone.now() - timedelta(minutes=1)
        
        # 查询新告警
        new_alerts = Alert.objects.filter(
            timestamp__gt=since,
            processed=False
        ).order_by('-timestamp')
        
        # 格式化告警数据
        alerts_data = []
        for alert in new_alerts:
            alerts_data.append({
                'id': alert.id,
                'elder_name': alert.elder.name,
                'device_name': alert.device.name if alert.device else None,
                'alert_type': alert.alert_type,
                'alert_level': alert.alert_level,
                'alert_level_display': alert.get_alert_level_display(),
                'description': alert.description,
                'timestamp': alert.timestamp.isoformat(),
                'location': alert.device.location if alert.device else '未知'
            })
        
        return JsonResponse({
            'status': 'success',
            'alerts': alerts_data
        })
    
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

@csrf_exempt
@require_POST
def process_alert(request):
    """API端点：处理告警"""
    try:
        print(f"收到处理告警请求: {request.body}")
        
        data = json.loads(request.body)
        alert_id = data.get('alert_id')
        
        print(f"处理告警ID: {alert_id}")
        
        if not alert_id:
            print("错误: 缺少告警ID")
            return JsonResponse({
                'status': 'error',
                'message': '缺少告警ID'
            }, status=400)
        
        # 获取告警
        try:
            alert = get_object_or_404(Alert, id=alert_id)
            print(f"找到告警: ID={alert.id}, 类型={alert.alert_type}")
        except Exception as e:
            print(f"获取告警失败: {str(e)}")
            return JsonResponse({
                'status': 'error',
                'message': f'获取告警失败: {str(e)}'
            }, status=404)
        
        # 更新告警状态
        try:
            alert.processed = True
            # 如果用户已登录，记录处理人
            if hasattr(request, 'user') and request.user.is_authenticated:
                alert.processed_by = request.user
                print(f"处理人: {request.user.username}")
            else:
                print("未登录用户处理告警")
            
            alert.processed_time = timezone.now()
            alert.save()
            print(f"告警已更新为已处理状态")
        except Exception as e:
            print(f"更新告警状态失败: {str(e)}")
            return JsonResponse({
                'status': 'error',
                'message': f'更新告警状态失败: {str(e)}'
            }, status=500)
        
        # 返回成功响应
        response_data = {
            'status': 'success',
            'message': '告警已处理',
            'alert_id': alert.id
        }
        print(f"返回响应: {response_data}")
        return JsonResponse(response_data)
    
    except json.JSONDecodeError as e:
        print(f"JSON解析错误: {str(e)}")
        return JsonResponse({
            'status': 'error',
            'message': f'JSON解析错误: {str(e)}'
        }, status=400)
    except Exception as e:
        print(f"处理告警请求时出错: {str(e)}")
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

# 添加一个API端点用于老人端触发紧急求助
@csrf_exempt
@require_POST
def trigger_emergency(request):
    """API端点：老人端触发紧急求助"""
    try:
        # 打印请求信息，便于调试
        print(f"收到紧急求助请求: {request.body}")
        
        data = json.loads(request.body)
        elder_id = data.get('elder_id')
        device_id = data.get('device_id')
        alert_type = data.get('alert_type', '紧急求助')
        description = data.get('description', '老人触发紧急求助')
        
        print(f"解析数据: elder_id={elder_id}, alert_type={alert_type}, description={description}")
        
        if not elder_id:
            print("错误: 缺少老人ID")
            return JsonResponse({
                'status': 'error',
                'message': '缺少老人ID'
            }, status=400)
        
        # 获取老人信息
        try:
            elder = get_object_or_404(Elder, id=elder_id)
            print(f"找到老人: {elder.name}")
        except Exception as e:
            print(f"获取老人信息失败: {str(e)}")
            return JsonResponse({
                'status': 'error',
                'message': f'获取老人信息失败: {str(e)}'
            }, status=400)
        
        # 获取设备信息（如果有）
        device = None
        if device_id:
            try:
                device = get_object_or_404(Device, device_id=device_id)
                print(f"找到设备: {device.name}")
            except Exception as e:
                print(f"获取设备信息失败: {str(e)}")
                # 设备不存在不影响主流程，继续执行
        
        # 创建紧急告警
        try:
            alert = Alert.objects.create(
                elder=elder,
                device=device,
                alert_type=alert_type,
                alert_level='critical',  # 紧急告警级别
                description=description,
                timestamp=timezone.now(),
                processed=False
            )
            print(f"成功创建告警: ID={alert.id}")
        except Exception as e:
            print(f"创建告警失败: {str(e)}")
            return JsonResponse({
                'status': 'error',
                'message': f'创建告警失败: {str(e)}'
            }, status=500)
        
        # 返回成功响应
        response_data = {
            'status': 'success',
            'message': '紧急求助已触发',
            'alert_id': alert.id
        }
        print(f"返回响应: {response_data}")
        return JsonResponse(response_data)
    
    except json.JSONDecodeError as e:
        print(f"JSON解析错误: {str(e)}")
        return JsonResponse({
            'status': 'error',
            'message': f'JSON解析错误: {str(e)}'
        }, status=400)
    except Exception as e:
        print(f"处理紧急求助请求时出错: {str(e)}")
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)
